#!/usr/bin/env python3
# make_reference.py -- build a themed pandoc reference.docx: the Onyx visual signature.
#
# Palette derived from the OS itself: onyx black + steel blue (Voronoi wallpaper) +
# gold accent (the active-window chrome tint). Used by docs/build_docs.py via
# `--reference-doc`. Run once (re-run if you change the theme):
#     python docs/assets/make_reference.py
#
import os, subprocess
import pypandoc
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
REF  = os.path.join(HERE, "reference.docx")

# ---- Onyx signature palette ----
ONYX  = RGBColor(0x10, 0x14, 0x17)   # near-black (titles)
INK   = RGBColor(0x22, 0x26, 0x2B)   # body text
NAVY  = RGBColor(0x16, 0x32, 0x4F)   # H1 / H3
STEEL = RGBColor(0x2F, 0x5C, 0x8F)   # H2 / links
SLATE = RGBColor(0x5A, 0x6B, 0x7B)   # H4 / footer
GOLD  = RGBColor(0xC8, 0x96, 0x2E)   # accent / rules / subtitle
CODE  = RGBColor(0x1B, 0x3A, 0x57)   # code text
GOLDX = "C8962E"; STEELX = "2F5C8F"; CODEFILL = "F2F4F7"; GRID = "C9D2DC"


def get_default_reference():
    pandoc = pypandoc.get_pandoc_path()
    data = subprocess.run([pandoc, "--print-default-data-file", "reference.docx"],
                          capture_output=True).stdout
    with open(REF, "wb") as f:
        f.write(data)


def font(style, name=None, size=None, color=None, bold=None, italic=None):
    f = style.font
    if name is not None:  f.name = name
    if size is not None:  f.size = Pt(size)
    if color is not None: f.color.rgb = color
    if bold is not None:  f.bold = bold
    if italic is not None: f.italic = italic


def bottom_border(style, color, sz=14, space=4):
    pPr = style.element.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), str(space)); b.set(qn('w:color'), color)
    pbdr.append(b)
    pPr.append(pbdr)


def para_shading(style, fill):
    pPr = style.element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def run_shading(style, fill):
    rPr = style.element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    rPr.append(shd)


def table_borders(style, color):
    el = style.element
    tblPr = el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); el.append(tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def add_page_field(par):
    par.add_run("page ")
    r = par.add_run()
    for t in ('begin', None, 'end'):
        if t is None:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
            r._r.append(it)
        else:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), t); r._r.append(fc)


def main():
    get_default_reference()
    doc = Document(REF)
    S = doc.styles

    def has(n):
        try: return S[n] is not None
        except KeyError: return False

    font(S['Normal'], 'Segoe UI', 10.5, INK)

    if has('Title'):
        font(S['Title'], 'Segoe UI Semilight', 30, ONYX, bold=True)
        bottom_border(S['Title'], GOLDX, sz=18, space=6)
    if has('Subtitle'):
        font(S['Subtitle'], 'Segoe UI', 13, GOLD, italic=True)

    if has('Heading 1'):
        font(S['Heading 1'], 'Segoe UI Semibold', 18, NAVY, bold=True)
        bottom_border(S['Heading 1'], STEELX, sz=8, space=4)
    if has('Heading 2'): font(S['Heading 2'], 'Segoe UI Semibold', 14, STEEL, bold=True)
    if has('Heading 3'): font(S['Heading 3'], 'Segoe UI Semibold', 12, NAVY, bold=True)
    if has('Heading 4'): font(S['Heading 4'], 'Segoe UI', 11, SLATE, bold=True, italic=True)

    if has('Hyperlink'): font(S['Hyperlink'], color=STEEL)

    for nm in ('Source Code', 'Verbatim Char'):
        if has(nm):
            font(S[nm], 'Consolas', 9.5, CODE)
            (para_shading if nm == 'Source Code' else run_shading)(S[nm], CODEFILL)

    if has('Table'):
        table_borders(S['Table'], GRID)

    # Page geometry + footer (project signature on every page)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)
    foot = sec.footer
    p = foot.paragraphs[0] if foot.paragraphs else foot.add_paragraph()
    p.text = ""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Onyx  ·  multi-process OS on Raspberry Pi 4      ")
    r.font.name = 'Segoe UI'; r.font.size = Pt(8.5); r.font.color.rgb = SLATE
    add_page_field(p)
    for rr in p.runs:
        rr.font.name = 'Segoe UI'; rr.font.size = Pt(8.5); rr.font.color.rgb = SLATE

    doc.save(REF)
    print("themed reference ->", REF)


if __name__ == "__main__":
    main()
